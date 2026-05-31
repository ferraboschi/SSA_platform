#!/usr/bin/env python3
"""
RAG ingestion — scarica i materiali del corso da Dropbox, estrae il testo,
lo divide in chunk, calcola gli embeddings (OpenAI) e li salva in Supabase
(rag_documents + rag_chunks con pgvector).

Tutte le credenziali sono lette da .env.local (mai hardcoded).
Dipendenze: python-docx, pypdf  →  python3 -m pip install python-docx pypdf

Uso:  python3 scripts/rag-ingest.py
"""
import json, os, re, time, urllib.request, urllib.error

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ── Config da .env.local ──────────────────────────────────────────────────────
env = {}
with open(os.path.join(ROOT, ".env.local")) as f:
    for line in f:
        line = line.strip()
        if "=" in line and not line.startswith("#"):
            k, v = line.split("=", 1)
            env[k] = v

SUPABASE_URL = env["NEXT_PUBLIC_SUPABASE_URL"]
SERVICE_KEY  = env["SUPABASE_SERVICE_ROLE_KEY"]
OPENAI_KEY   = env["EMBEDDINGS_API_KEY"]
DBX_TOKEN    = env["DROPBOX_ACCESS_TOKEN"]
MODEL        = env.get("EMBEDDINGS_MODEL", "text-embedding-3-small")

SB_H = {"apikey": SERVICE_KEY, "Authorization": f"Bearer {SERVICE_KEY}",
        "Content-Type": "application/json", "Prefer": "return=representation"}

def sb(method, path, data=None, extra=None):
    url = f"{SUPABASE_URL}/rest/v1/{path}"
    body = json.dumps(data).encode() if data is not None else None
    req = urllib.request.Request(url, data=body,
        headers={**SB_H, **(extra or {})}, method=method)
    try:
        with urllib.request.urlopen(req) as r:
            txt = r.read()
            return (json.loads(txt) if txt else None), None
    except urllib.error.HTTPError as e:
        return None, e.read().decode()[:300]

def embed(texts):
    body = json.dumps({"model": MODEL, "input": texts}).encode()
    req = urllib.request.Request("https://api.openai.com/v1/embeddings",
        data=body, headers={"Authorization": f"Bearer {OPENAI_KEY}",
        "Content-Type": "application/json"}, method="POST")
    with urllib.request.urlopen(req) as r:
        return [d["embedding"] for d in json.loads(r.read())["data"]]

# ── Download da Dropbox ───────────────────────────────────────────────────────
SHARED_LINK = ("https://www.dropbox.com/scl/fo/6v3emocwmddcvc1jvbhw8/"
               "ABx2TS9OTni6H3-AbpOVOyA?rlkey=x1gcjz1ml1tu0h5puijjbui7j&dl=0")

def dbx_download_path(dbx_path, dest):
    arg = json.dumps({"path": dbx_path})
    req = urllib.request.Request("https://content.dropboxapi.com/2/files/download",
        headers={"Authorization": f"Bearer {DBX_TOKEN}", "Dropbox-API-Arg": arg},
        method="POST")
    with urllib.request.urlopen(req) as r:
        open(dest, "wb").write(r.read())

def dbx_download_shared(path_in_folder, dest):
    arg = json.dumps({"url": SHARED_LINK, "path": "/" + path_in_folder})
    req = urllib.request.Request(
        "https://content.dropboxapi.com/2/sharing/get_shared_link_file",
        headers={"Authorization": f"Bearer {DBX_TOKEN}", "Dropbox-API-Arg": arg},
        method="POST")
    with urllib.request.urlopen(req) as r:
        open(dest, "wb").write(r.read())

# ── Estrazione testo ──────────────────────────────────────────────────────────
from docx import Document
def docx_text(path):
    doc = Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells: parts.append(" | ".join(cells))
    return "\n".join(parts)

# ── Chunking ──────────────────────────────────────────────────────────────────
def chunk_text(text, target=1400, overlap=200):
    paras = [p.strip() for p in text.split("\n") if p.strip()]
    chunks, cur = [], ""
    for p in paras:
        while len(p) > target * 1.5:
            chunks.append(p[:target]); p = p[target - overlap:]
        if len(cur) + len(p) + 1 <= target:
            cur += (("\n" if cur else "") + p)
        else:
            if cur: chunks.append(cur)
            tail = chunks[-1][-overlap:] if (overlap and chunks) else ""
            cur = (tail + "\n" + p) if tail else p
    if cur.strip(): chunks.append(cur)
    return [c.strip() for c in chunks if len(c.strip()) > 30]

# ── Corpus ────────────────────────────────────────────────────────────────────
# (source, titolo, family, modo-download, path)
TMP = "/tmp/ssa-rag"
os.makedirs(TMP, exist_ok=True)
CORPUS = [
    ("path",   "/SSA/Sussidiario/TESTI DISPENSA 2023.docx",       "Dispensa SSA 2023 (IT)",           "generale"),
    ("path",   "/SSA/Sussidiario/TESTI DISPENSA 2023 en.docx",    "SSA Handout 2023 (EN)",            "generale"),
    ("path",   "/SSA/Handouts/GLOSSARIO.docx",                    "Glossario SSA",                    "generale"),
    ("shared", "Introduzione_OK.docx",                            "Libro · Introduzione",             "generale"),
    ("shared", "Capitolo 1 - Il sake_07_23.docx",                 "Libro · Cap 1 — Il sake",          "generale"),
    ("shared", "Capitolo 2 - Fare il sake_7-9.docx",              "Libro · Cap 2 — Fare il sake",     "generale"),
    ("shared", "Capitolo 3 - Tipologie_bozza.docx",               "Libro · Cap 3 — Tipologie",        "generale"),
    ("shared", "Capitolo 3 - Come bere il sake_7_30.docx",        "Libro · Cap 3 — Come bere il sake","generale"),
    ("shared", "Capitolo 4 - Degustazione_bozza.docx",            "Libro · Cap 4 — Degustazione",     "generale"),
    ("shared", "Approfondimento sulla Noto Toji_7_30.docx",       "Libro · Approfondimento Noto Toji","generale"),
]

def main():
    print("Pulisco RAG esistente…")
    sb("DELETE", "rag_documents?id=gt.0")
    total = 0
    for mode, src, title, family in CORPUS:
        fname = src.split("/")[-1]
        dest = os.path.join(TMP, fname)
        try:
            (dbx_download_path if mode == "path" else dbx_download_shared)(src, dest)
        except Exception as e:
            print(f"  ❌ download {title}: {e}"); continue
        chunks = chunk_text(docx_text(dest))
        if not chunks:
            print(f"  ⚠️  {title}: vuoto"); continue
        doc, err = sb("POST", "rag_documents", [{
            "source": f"dropbox:{src}", "title": title, "family": family,
            "metadata": {"file": fname}}])
        if err:
            print(f"  ❌ {title}: {err}"); continue
        doc_id = doc[0]["id"]
        rows = []
        for i in range(0, len(chunks), 96):
            batch = chunks[i:i+96]
            for j, (c, v) in enumerate(zip(batch, embed(batch))):
                rows.append({"document_id": doc_id, "chunk_index": i + j,
                    "content": c,
                    "embedding": "[" + ",".join(f"{x:.6f}" for x in v) + "]",
                    "metadata": {"title": title}})
            time.sleep(0.2)
        for i in range(0, len(rows), 50):
            sb("POST", "rag_chunks", rows[i:i+50], {"Prefer": "return=minimal"})
        total += len(rows)
        print(f"  ✅ {title}: {len(rows)} chunk")
    print(f"\nTotale: {total} chunk indicizzati.")

if __name__ == "__main__":
    main()
